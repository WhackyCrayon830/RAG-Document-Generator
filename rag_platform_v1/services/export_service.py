"""Export service - converts markdown docs to PDF, DOCX, HTML."""
import os
import logging

logger = logging.getLogger(__name__)


def export_to_html(markdown_text: str, output_path: str) -> bool:
    try:
        import markdown as md_lib
        html_content = md_lib.markdown(
            markdown_text,
            extensions=["tables", "fenced_code", "toc"]
        )
        full_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
  body {{ font-family: Arial, sans-serif; max-width: 900px; margin: 40px auto; padding: 0 20px; }}
  table {{ border-collapse: collapse; width: 100%; }}
  th, td {{ border: 1px solid #ddd; padding: 8px; }}
  pre {{ background: #f4f4f4; padding: 12px; border-radius: 4px; overflow-x: auto; }}
  code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
  blockquote {{ border-left: 4px solid #ccc; margin: 0; padding-left: 16px; color: #666; }}
</style>
</head>
<body>{html_content}</body>
</html>"""
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_html)
        return True
    except Exception as e:
        logger.error(f"HTML export failed: {e}")
        return False


def export_to_docx(markdown_text: str, output_path: str) -> bool:
    try:
        from docx import Document
        from docx.shared import Pt
        import re

        doc = Document()
        lines = markdown_text.split("\n")
        for line in lines:
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("```"):
                pass  # skip fence markers
            elif line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph("")

        doc.save(output_path)
        return True
    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        return False


def export_to_pdf(markdown_text: str, output_path: str) -> bool:
    """Export markdown to PDF via HTML intermediate."""
    try:
        html_tmp = output_path.replace(".pdf", "_tmp.html")
        export_to_html(markdown_text, html_tmp)

        try:
            import weasyprint
            weasyprint.HTML(filename=html_tmp).write_pdf(output_path)
            os.remove(html_tmp)
            return True
        except ImportError:
            # Fallback: use reportlab for basic PDF
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                doc = SimpleDocTemplate(output_path, pagesize=A4)
                styles = getSampleStyleSheet()
                story = []
                for line in markdown_text.split("\n"):
                    clean = line.strip().lstrip("#").strip()
                    if clean:
                        story.append(Paragraph(clean, styles["Normal"]))
                        story.append(Spacer(1, 6))
                doc.build(story)
                if os.path.exists(html_tmp):
                    os.remove(html_tmp)
                return True
            except ImportError:
                logger.warning("No PDF library available. Install weasyprint or reportlab.")
                return False

    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        return False
