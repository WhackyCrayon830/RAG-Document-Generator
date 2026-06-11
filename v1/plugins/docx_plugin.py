"""DOCX processor plugin - preserves heading hierarchy and formatting."""
import os
from typing import List
from plugins.base_plugin import BasePlugin, DocumentChunk


class DOCXPlugin(BasePlugin):
    SUPPORTED_EXTENSIONS = [".docx"]

    def extract(self, file_path: str) -> List[DocumentChunk]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        if not self.validate_file(file_path):
            return []

        chunks: List[DocumentChunk] = []
        source = os.path.basename(file_path)
        current_section = ""
        buffer: List[str] = []

        def flush_buffer():
            text = "\n".join(buffer).strip()
            if text:
                chunks.append(DocumentChunk(
                    text=text,
                    source=source,
                    page=0,
                    section=current_section,
                    chunk_type="text",
                ))
            buffer.clear()

        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                style = para.style.name if para.style else ""
                text = para.text.strip()
                if not text:
                    continue

                if "Heading" in style:
                    flush_buffer()
                    current_section = text
                    chunks.append(DocumentChunk(
                        text=text,
                        source=source,
                        page=0,
                        section=current_section,
                        chunk_type="heading",
                        metadata={"style": style},
                    ))
                else:
                    buffer.append(text)

            flush_buffer()

            # Extract tables
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if not rows:
                    continue
                headers = rows[0]
                lines = [" | ".join(headers), "-" * 40]
                for row in rows[1:]:
                    lines.append(" | ".join(row))
                chunks.append(DocumentChunk(
                    text="\n".join(lines),
                    source=source,
                    page=0,
                    section=current_section,
                    chunk_type="table",
                    metadata={"headers": headers},
                ))

        except Exception as e:
            chunks.append(DocumentChunk(
                text=f"[ERROR extracting DOCX: {e}]",
                source=source,
                page=0,
                chunk_type="error",
            ))

        return chunks
