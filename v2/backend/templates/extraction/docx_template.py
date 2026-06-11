from pathlib import Path

from docx import Document


def extract_template_profile(path: Path) -> dict:
    doc = Document(str(path))
    styles = []
    for style in doc.styles:
        if style.type:
            font = getattr(style, "font", None)
            styles.append(
                {
                    "name": style.name,
                    "font": getattr(font, "name", None) if font else None,
                    "bold": getattr(font, "bold", None) if font else None,
                    "italic": getattr(font, "italic", None) if font else None,
                }
            )
    placeholders = []
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "{{" in text and "}}" in text:
            placeholders.append(text)

    sections = []
    for section in doc.sections:
        sections.append(
            {
                "page_width": section.page_width.twips if section.page_width else None,
                "page_height": section.page_height.twips if section.page_height else None,
                "top_margin": section.top_margin.twips if section.top_margin else None,
                "bottom_margin": section.bottom_margin.twips if section.bottom_margin else None,
                "left_margin": section.left_margin.twips if section.left_margin else None,
                "right_margin": section.right_margin.twips if section.right_margin else None,
                "header": "\n".join(p.text for p in section.header.paragraphs if p.text.strip()),
                "footer": "\n".join(p.text for p in section.footer.paragraphs if p.text.strip()),
            }
        )

    numbering_styles = [
        style.name
        for style in doc.styles
        if "list" in style.name.lower() or "number" in style.name.lower()
    ]
    table_styles = [style.name for style in doc.styles if "table" in style.name.lower()]
    heading_levels = [
        {"style": style.name, "level": style.name.replace("Heading", "").strip()}
        for style in doc.styles
        if style.name.lower().startswith("heading")
    ]

    return {
        "styles": styles[:120],
        "placeholders": placeholders,
        "sections": sections,
        "heading_levels": heading_levels,
        "numbering_styles": numbering_styles,
        "table_styles": table_styles,
    }
