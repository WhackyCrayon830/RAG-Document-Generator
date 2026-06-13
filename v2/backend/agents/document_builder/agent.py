"""
Document Builder Agent – VLM Layout-Aware python-docx Code Generator

Flow:
  1. Inspect the template file (DOCX, PDF, or image) to extract layout/style details.
     - DOCX: read fonts, colours, margins, heading styles via python-docx.
     - PDF:  convert page 1 to PNG → pass to VLM; also extract text via pdfplumber.
     - Image (PNG/JPG): encode to base64 → pass to VLM.
  2. Build a style description string.
  3. Ask the writing LLM to generate a python-docx script that recreates the layout
     and injects the document sections.
  4. Execute the generated script (up to 3 self-correction attempts on failure).
  5. Fall back to the basic compile_docx helper if all attempts fail.
"""

from __future__ import annotations

import base64
import logging
import subprocess
import sys
import tempfile
import textwrap
from io import BytesIO
from pathlib import Path

from backend.agents.ollama_client import OllamaClient
from backend.config.settings import Settings
from backend.exporters.docx.compiler import compile_docx

logger = logging.getLogger(__name__)

_MAX_SCRIPT_ATTEMPTS = 3


class DocumentBuilderAgent:
    """Generate a styled DOCX by having the LLM write and execute python-docx code."""

    def __init__(
        self,
        ollama: OllamaClient,
        settings: Settings,
        vision_model: str | None = None,
        writing_model: str | None = None,
    ):
        self.ollama = ollama
        self.settings = settings
        self.vision_model = vision_model or settings.ollama_vision_model
        self.writing_model = writing_model or settings.ollama_writing_model

    # ─────────────────────────── Public entry point ───────────────────────────

    def build(
        self,
        title: str,
        sections: list[dict],
        output_path: Path,
        template_path: Path | None = None,
    ) -> Path:
        """Build the final DOCX. Returns path to the generated file."""
        output_path.parent.mkdir(parents=True, exist_ok=True)

        style_description = self._extract_style_description(template_path)
        sections_text = self._sections_to_text(sections)

        script = self._generate_script(title, sections_text, style_description, str(output_path))

        for attempt in range(1, _MAX_SCRIPT_ATTEMPTS + 1):
            error = self._execute_script(script, output_path)
            if error is None:
                logger.info("DocumentBuilderAgent: script succeeded on attempt %d", attempt)
                return output_path
            logger.warning("DocumentBuilderAgent: attempt %d failed: %s", attempt, error[:300])
            if attempt < _MAX_SCRIPT_ATTEMPTS:
                script = self._fix_script(script, error)

        # All script attempts failed – fall back to basic compiler
        logger.error("DocumentBuilderAgent: all script attempts failed; using basic compiler.")
        return compile_docx(title, sections, output_path, template_path=template_path)

    # ─────────────────────────── Template inspection ─────────────────────────

    def _extract_style_description(self, template_path: Path | None) -> str:
        if template_path is None or not template_path.exists():
            return "No template provided. Use clean, professional styling with headings and body paragraphs."

        suffix = template_path.suffix.lower()

        if suffix == ".docx":
            return self._describe_docx(template_path)
        elif suffix == ".pdf":
            return self._describe_pdf(template_path)
        elif suffix in (".png", ".jpg", ".jpeg", ".bmp", ".webp"):
            return self._describe_image(template_path)
        else:
            return f"Template file type '{suffix}' not directly supported for style extraction. Use clean default styling."

    def _describe_docx(self, path: Path) -> str:
        """Extract style description from a DOCX template using python-docx."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            doc = Document(str(path))

            lines: list[str] = ["DOCX Template Style Analysis:"]

            # Margins
            for section in doc.sections:
                lines.append(
                    f"Margins: top={section.top_margin.inches:.2f}in, "
                    f"bottom={section.bottom_margin.inches:.2f}in, "
                    f"left={section.left_margin.inches:.2f}in, "
                    f"right={section.right_margin.inches:.2f}in"
                )
                break

            # Styles
            for style in doc.styles:
                if style.font and style.name in ("Normal", "Heading 1", "Heading 2", "Title"):
                    font = style.font
                    color_str = ""
                    if font.color and font.color.rgb:
                        color_str = f", color=#{font.color.rgb}"
                    lines.append(
                        f"Style '{style.name}': font={font.name or 'default'}, "
                        f"size={font.size.pt if font.size else 'default'}pt"
                        f"{color_str}"
                    )

            return "\n".join(lines)
        except Exception as exc:
            logger.warning("DOCX style extraction failed: %s", exc)
            return "Professional DOCX template detected. Use clean heading and body paragraph styles."

    def _describe_pdf(self, path: Path) -> str:
        """Extract style description from PDF using PyMuPDF (page→PNG→VLM) and pdfplumber."""
        text_sample = ""
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                if pdf.pages:
                    text_sample = (pdf.pages[0].extract_text() or "")[:500]
        except Exception as exc:
            logger.warning("pdfplumber extraction failed: %s", exc)

        # Attempt VLM image analysis
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            page = doc[0]
            pix = page.get_pixmap(dpi=96)
            img_bytes = pix.tobytes("png")
            b64 = base64.b64encode(img_bytes).decode()
            doc.close()

            vlm_response = self.ollama.generate(
                model=self.vision_model,
                prompt=(
                    "Describe the layout and visual style of this document page in detail. "
                    "Include: font families, approximate font sizes, colours, margins, "
                    "header/footer presence, column layout, logo placement, and any notable design elements. "
                    "Be concise and specific."
                ),
                images=[b64],
            )
            if vlm_response:
                description = f"PDF Template VLM Analysis:\n{vlm_response}"
                if text_sample:
                    description += f"\n\nSample text from page 1:\n{text_sample}"
                return description
        except Exception as exc:
            logger.warning("PDF VLM analysis failed: %s", exc)

        return f"PDF template detected. Sample text:\n{text_sample}\nUse professional document styling."

    def _describe_image(self, path: Path) -> str:
        """Describe an image template via VLM."""
        try:
            b64 = base64.b64encode(path.read_bytes()).decode()
            vlm_response = self.ollama.generate(
                model=self.vision_model,
                prompt=(
                    "Describe the layout and visual style of this document template image in detail. "
                    "Include: font families, colours, margins, logo/header placement, "
                    "section structure, and any notable design elements."
                ),
                images=[b64],
            )
            if vlm_response:
                return f"Image Template VLM Analysis:\n{vlm_response}"
        except Exception as exc:
            logger.warning("Image VLM analysis failed: %s", exc)

        return "Image template detected. Use clean professional styling."

    # ─────────────────────────── Script generation ────────────────────────────

    def _sections_to_text(self, sections: list[dict]) -> str:
        """Render sections as a structured text block for injection into LLM prompt."""
        parts: list[str] = []
        for sec in sections:
            parts.append(f"## {sec['title']}\n{sec.get('content', '').strip()}")
        return "\n\n".join(parts)

    def _generate_script(
        self,
        title: str,
        sections_text: str,
        style_description: str,
        output_path_str: str,
    ) -> str:
        """Ask the LLM to generate a runnable python-docx script."""
        # Keep sections_text short to avoid LLM context overflow
        truncated_sections = sections_text[:6000] + ("\n... [truncated]" if len(sections_text) > 6000 else "")

        prompt = textwrap.dedent(f"""
        Write a complete Python script using the python-docx library to create a professional document.

        TEMPLATE STYLE DESCRIPTION:
        {style_description}

        DOCUMENT TITLE: {title}

        SECTIONS (inject this content exactly):
        {truncated_sections}

        REQUIREMENTS:
        - Use `from docx import Document` and `from docx.shared import Pt, RGBColor, Inches, Cm` as needed.
        - Save the document to: {output_path_str!r}
        - Apply the template styles described above as closely as possible.
        - For each section, add a Heading 1 with the section title, then add paragraphs for the content.
        - Split section content on double-newlines to create separate paragraphs.
        - Do NOT add a cover page or table of contents unless the template explicitly requires it.
        - The script must be self-contained and executable with no user input.
        - Output ONLY the Python code. Do not include markdown code fences, explanations, or comments outside the code.
        """).strip()

        response = self.ollama.generate(
            model=self.writing_model,
            prompt=prompt,
            system=(
                "You are a Python code generation agent. "
                "Output only valid, runnable Python code. "
                "No markdown. No explanations. No ```python fences."
            ),
        )

        if not response:
            # Minimal fallback script
            return self._minimal_script(title, sections_text, output_path_str)

        # Strip any accidental markdown fences
        code = response.strip()
        if code.startswith("```"):
            lines = code.splitlines()
            code = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])

        return code

    def _fix_script(self, script: str, error: str) -> str:
        """Ask the LLM to fix a broken script given the traceback."""
        prompt = textwrap.dedent(f"""
        The following Python script using python-docx raised an error when executed.
        Fix the script so it runs correctly.

        ERROR:
        {error[:1200]}

        ORIGINAL SCRIPT:
        {script[:4000]}

        Output ONLY the corrected Python code. No explanations, no markdown fences.
        """).strip()

        response = self.ollama.generate(
            model=self.writing_model,
            prompt=prompt,
            system="You are a Python debugging agent. Output only corrected Python code.",
        )

        if not response:
            return script

        code = response.strip()
        if code.startswith("```"):
            lines = code.splitlines()
            code = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
        return code

    # ─────────────────────────── Script execution ─────────────────────────────

    def _execute_script(self, script: str, output_path: Path) -> str | None:
        """Execute the generated script in a subprocess. Returns None on success, error string on failure."""
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".py", encoding="utf-8", delete=False
        ) as fh:
            fh.write(script)
            tmp_path = fh.name

        try:
            result = subprocess.run(
                [sys.executable, tmp_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0 and output_path.exists():
                return None
            error = (result.stderr or result.stdout or "Unknown error").strip()
            return error
        except subprocess.TimeoutExpired:
            return "Script execution timed out after 60 seconds."
        except Exception as exc:
            return str(exc)
        finally:
            try:
                Path(tmp_path).unlink()
            except Exception:
                pass

    # ─────────────────────────── Minimal fallback script ─────────────────────

    def _minimal_script(self, title: str, sections_text: str, output_path_str: str) -> str:
        """Return a simple but reliable python-docx script as last resort."""
        # Escape for embedding in source
        safe_title = title.replace("'", "\\'")
        safe_out = output_path_str.replace("\\", "\\\\").replace("'", "\\'")
        sections_repr = repr(sections_text[:4000])

        return textwrap.dedent(f"""
from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('{safe_title}', level=0)

sections_text = {sections_repr}
current_heading = None
for line in sections_text.splitlines():
    stripped = line.strip()
    if stripped.startswith('## '):
        current_heading = stripped[3:]
        doc.add_heading(current_heading, level=1)
    elif stripped:
        doc.add_paragraph(stripped)

doc.save('{safe_out}')
""").strip()
