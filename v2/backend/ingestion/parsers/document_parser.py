from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def parse_document(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if suffix == ".docx":
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append(" | ".join(cell.text.strip() for cell in row.cells))
            if rows:
                tables.append("\n".join(rows))
        return "\n\n".join(paragraphs + tables).strip()
    raise ValueError(f"Unsupported file type: {suffix}")
